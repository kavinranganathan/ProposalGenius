import streamlit as st
from groq import Groq
import PyPDF2
from docx import Document
import io
import os
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Initialize GROQ client
api_key = os.getenv("GROQ_API_KEY")  # Read API key from environment variable
if not api_key:
    st.error("GROQ_API_KEY not found in environment variables. Please add it to your .env file.")
    st.stop()

client = Groq(api_key=api_key)

# Function to extract text from PDF
def extract_text_from_pdf(pdf_file):
    pdf_reader = PyPDF2.PdfReader(pdf_file)  # Use PdfReader instead of PdfFileReader
    text = ""
    for page_num in range(len(pdf_reader.pages)):  # Use len(pdf_reader.pages) instead of numPages
        page = pdf_reader.pages[page_num]  # Use .pages instead of .getPage()
        text += page.extract_text()
    return text

# Function to extract text from DOCX
def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

# Function to create a Word document
def create_word_document(text):
    doc = Document()
    doc.add_paragraph(text)
    return doc

# Function to preprocess text (remove unnecessary sections)
def preprocess_text(text):
    # Example: Remove headers, footers, or boilerplate text
    lines = text.split("\n")
    cleaned_lines = [line for line in lines if line.strip() and not line.startswith("Page")]
    return "\n".join(cleaned_lines)

# Streamlit App
st.title("ProposalGenius - Intelligent RFP Proposal Maker")

# File upload
uploaded_file = st.file_uploader("Upload a PDF or DOCX file", type=["pdf", "docx"])

if uploaded_file is not None:
    file_details = {"filename": uploaded_file.name, "filetype": uploaded_file.type, "filesize": uploaded_file.size}
    st.write(file_details)

    # Extract text based on file type
    if uploaded_file.type == "application/pdf":
        text = extract_text_from_pdf(uploaded_file)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        text = extract_text_from_docx(uploaded_file)
    else:
        st.error("Unsupported file type. Please upload a PDF or DOCX file.")
        st.stop()

    # Preprocess the text
    cleaned_text = preprocess_text(text)

    # Display extracted text
    st.subheader("Extracted Text")
    st.text_area("Text", cleaned_text, height=300)

    # Truncate the input text to avoid exceeding token limits
    max_tokens = 4000  # Adjust this value based on the model's token limit
    truncated_text = cleaned_text[:max_tokens]

    # GROQ Chat Completion with Streaming
    st.subheader("Generate Proposal with ProposalGenius")
    if st.button("Generate Proposal"):
        with st.spinner("Generating proposal..."):  # Show a spinner while processing
            # Stream the response from GROQ
            stream = client.chat.completions.create(
                messages=[
                    {
                        "role": "system",
                        "content": """
                        You are a professional proposal writer. Generate a detailed, accurate, and professional RFP proposal based on the provided text.
                        The proposal must follow this structure:
                        
                        1. **Cover Letter**
                           - Brief introduction of your organization.
                           - Acknowledgment of the RFP and the intent to provide a solution.
                           - High-level summary of your qualifications and approach.
                           - Contact information for follow-up.

                        2. **Executive Summary**
                           - Overview of the proposal.
                           - Key features of the proposed solution or service.
                           - Unique selling points (why you are the best fit).
                           - Summary of benefits to the client.

                        3. **Understanding of the Client’s Requirements**
                           - Demonstrate a clear understanding of the RFP’s objectives and scope.
                           - Highlight specific needs or challenges addressed in the RFP.
                           - Show alignment with the client’s goals.

                        4. **Proposed Solution/Approach**
                           - Detailed description of your solution or service offering.
                           - Methodology, processes, or approach to deliver the requirements.
                           - Innovative or tailored aspects of the solution.
                           - Timeline or milestones for the project.

                        5. **Organization Overview and Credentials**
                           - Company history and background.
                           - Relevant qualifications, certifications, or licenses.
                           - Key team members and their expertise.
                           - Case studies or examples of similar projects successfully completed.

                        6. **Technical and Functional Specifications**
                           - Detailed technical features or functional aspects of the solution.
                           - Compatibility with the client’s current systems or workflows.
                           - Scalability, security, and other technical considerations.

                        7. **Implementation Plan**
                           - Project timeline, including key phases and milestones.
                           - Resource allocation and responsibilities.
                           - Training, support, or change management offered.
                           - Risk management and mitigation strategies.

                        8. **Cost Proposal**
                           - Detailed pricing breakdown (e.g., one-time costs, recurring costs, optional features).
                           - Payment terms and conditions.
                           - Cost-saving advantages or added value.

                        9. **Service Levels and Support**
                           - Service Level Agreements (SLAs).
                           - Ongoing support, maintenance, or warranty details.
                           - Escalation procedures for issues.

                        10. **Compliance and Certifications**
                            - Adherence to any regulatory, industry, or security standards specified in the RFP.
                            - Relevant certifications or audits (e.g., ISO, SOC, GDPR compliance).

                        11. **Value Proposition and Differentiators**
                            - What sets your organization apart from competitors.
                            - Unique features, innovations, or approaches.
                            - Long-term benefits and ROI for the client.

                        12. **Appendices**
                            - Supporting documents (e.g., resumes, case studies, references, technical specifications).
                            - Terms and conditions or legal disclaimers.
                            - Additional information requested in the RFP (e.g., forms, checklists).

                        13. **Signature Page**
                            - A formal section for authorized signatories to confirm the proposal.
                            - Include the date and legal acknowledgment of the proposal.

                        Ensure the proposal is well-structured, clear, and tailored to the client's requirements.
                        """
                    },
                    {
                        "role": "user",
                        "content": f"Generate a detailed and accurate RFP proposal based on the following text: {truncated_text}"
                    }
                ],
                model="llama-3.3-70b-versatile",  # Use a model with a larger context window
                temperature=0.5,
                max_tokens=2048,  # Increase tokens for longer proposals
                top_p=1,
                stop=None,
                stream=True,  # Enable streaming
            )

            # Display streaming response
            st.write("ProposalGenius is generating your proposal...")
            response_container = st.empty()
            full_response = ""
            for chunk in stream:
                if chunk.choices[0].delta.content:
                    full_response += chunk.choices[0].delta.content
                    response_container.markdown(full_response)

            # Create Word document with the generated proposal
            doc = create_word_document(full_response)

            # Save Word document to a BytesIO object
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)

            # Download button for Word document
            st.download_button(
                label="Download Proposal",
                data=doc_bytes,
                file_name="generated_proposal.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
